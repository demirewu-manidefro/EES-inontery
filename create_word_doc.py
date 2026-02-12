from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def create_word_doc():
    doc = Document()
    
    # --- Title Page ---
    doc.add_section()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Employee Material Management System\n(EMMS)")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("\nProject Documentation & Technical Reference\n")
    run.font.size = Pt(16)
    run.italic = True
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"\n\nPrepared for: Ethiopian Statistics Service (ESS)\nDate: January 2026")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # --- 1. Introduction ---
    add_heading(doc, "1. Project Overview", 1)
    doc.add_paragraph("The Employee Material Management System (EMMS) is a centralized web-based platform designed to modernize asset tracking within the Ethiopian Statistics Service (ESS). It transitions the organization from manual, paper-based records to a secure, digital database.")
    
    add_heading(doc, "1.1 Problem Statement", 2)
    p = doc.add_paragraph()
    p.add_run("• Loss of Data:").bold = True
    p.add_run(" Paper registries are prone to damage and misplacement.\n")
    p.add_run("• Inefficiency:").bold = True
    p.add_run(" Searching for specific item history takes significant manual effort.\n")
    p.add_run("• Accountability Gap:").bold = True
    p.add_run(" Difficulty in tracking unreturned items to specific personnel.")

    add_heading(doc, "1.2 Core Objectives", 2)
    doc.add_paragraph("• Automate the borrowing and returning lifecycle of office materials.", style='List Bullet')
    doc.add_paragraph("• Ensure real-time visibility of inventory status (Available vs. Borrowed).", style='List Bullet')
    doc.add_paragraph("• Provide role-based access control for Admins and Standard Users.", style='List Bullet')

    # --- 2. Technical Architecture ---
    add_heading(doc, "2. Technical Architecture", 1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology Used'
    
    data = [
        ("Backend Framework", "Python (Flask 2.3)"),
        ("Database", "PostgreSQL (Production) / SQLite (Dev)"),
        ("ORM Layer", "SQLAlchemy"),
        ("Frontend UI", "HTML5, Bootstrap 5, Lucide Icons"),
        ("Authentication", "Flask-Login (Session Management)"),
        ("Reporting", "Pandas & OpenPyXL (Excel Exports)")
    ]
    
    for item, tech in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = tech
        
    doc.add_paragraph("\n")

    # --- 3. Functional Modules ---
    add_heading(doc, "3. Key Functional Modules", 1)
    
    add_heading(doc, "3.1 Workforce Management", 2)
    doc.add_paragraph("Managers can maintain a digital registry of all employees. Key data points captured include:")
    doc.add_paragraph("• Personal: Full Name (Father/Grandfather), Sex, Phone Number.", style='List Bullet')
    doc.add_paragraph("• Professional: Position, Employment Status (Contract/Permanent), Assigned Project.", style='List Bullet')
    
    add_heading(doc, "3.2 Inventory Control", 2)
    doc.add_paragraph("Assets are tracked individually using unique attributes:")
    doc.add_paragraph("• Material Name (e.g., 'Dell Latitude 5420')", style='List Bullet')
    doc.add_paragraph("• Serial Number (Unique Identifier)", style='List Bullet')
    doc.add_paragraph("• Status (Available / Borrowed / Maintenance)", style='List Bullet')
    
    add_heading(doc, "3.3 Deployment Logistics", 2)
    doc.add_paragraph("The system handles the full lifecycle of equipment lending:")
    doc.add_paragraph("1. Issue Material: Select an employee and assign multiple available items.")
    doc.add_paragraph("2. Return Process: Support for both bulk returns (clearance) and individual item returns.")
    doc.add_paragraph("3. Waitlist System: A dedicated queue for employees pending clearance.")

    # --- 4. Database Schema ---
    add_heading(doc, "4. Database Schema", 1)
    doc.add_paragraph("The application uses a relational database model with the following core entities:")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Table Name'
    hdr_cells[1].text = 'Description'
    
    schema_data = [
        ("employees", "Stores personnel details (Name, Sex, Phone, Project, Status)."),
        ("materials", "Inventory items with unique serial numbers and current status."),
        ("borrowed_materials", "Link table tracking which employee has which material and since when."),
        ("users", "System accounts (Admin/Manager) with hashed passwords and approval status."),
        ("leave_out_members", "Archive of employees who have left the organization."),
        ("waiting_return", "Temporary list for staff in the process of clearing their equipment.")
    ]
    
    for name, desc in schema_data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = desc

    doc.add_paragraph("\n")

    # --- 5. Installation ---
    add_heading(doc, "5. Installation & Setup", 1)
    doc.add_paragraph("To deploy the application locally:")
    doc.add_paragraph("1. Clone the repository and navigate to the project folder.")
    doc.add_paragraph("2. Install dependencies: pip install -r requirements.txt")
    doc.add_paragraph("3. Configure database URI in config.py or environment variables.")
    doc.add_paragraph("4. Initialize database: python init_db.py")
    doc.add_paragraph("5. Run server: flask run")

    doc.save("Employee_System_Project_Report.docx")
    print("Successfully created Employee_System_Project_Report.docx")

if __name__ == "__main__":
    create_word_doc()
